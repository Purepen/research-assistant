"""
Projects Routes — with Download Endpoint

Endpoint added: GET /projects/{project_id}/download
Returns a .docx file built from specification_json stored in the database.
"""

from __future__ import annotations

import io
from typing import Optional, List

from fastapi import APIRouter, Depends, HTTPException, Query, status
from fastapi.responses import StreamingResponse
from fastapi.security import HTTPBearer
from pydantic import BaseModel

from app.models.database import Project
from app.api.dependencies import get_db_session, get_current_user

router   = APIRouter(prefix="/projects", tags=["Projects"])
security = HTTPBearer()


# ── Response models ────────────────────────────────────────────────────────────

class ProjectListItem(BaseModel):
    id: int
    field_of_study: str
    research_topic: Optional[str]
    academic_level: str
    status: str
    progress_percentage: int
    created_at: str
    completed_at: Optional[str]
    total_marks: Optional[int]
    decision: Optional[str]


class ProjectDetail(BaseModel):
    id: int
    field_of_study: str
    research_topic: Optional[str]
    academic_level: str
    effort_level: str
    past_projects_mode: str
    status: str
    progress_percentage: int
    current_phase: Optional[str]
    created_at: str
    started_at: Optional[str]
    completed_at: Optional[str]
    has_results: bool
    total_marks: Optional[int]
    decision: Optional[str]


# ── List ───────────────────────────────────────────────────────────────────────

@router.get("/", response_model=List[ProjectListItem])
async def list_projects(
    skip:   int            = Query(0,  ge=0),
    limit:  int            = Query(20, ge=1, le=100),
    status: Optional[str]  = Query(None),
    user   = Depends(get_current_user),
    db     = Depends(get_db_session),
):
    query = db.query(Project).filter(Project.user_id == user.id)

    if status:
        from app.models.database import ProjectStatus
        try:
            status_enum = ProjectStatus[status.upper()]
            query = query.filter(Project.status == status_enum)
        except KeyError:
            pass

    projects = query.order_by(Project.created_at.desc()).offset(skip).limit(limit).all()

    return [
        {
            "id":                  p.id,
            "field_of_study":      p.field_of_study,
            "research_topic":      p.research_topic,
            "academic_level":      p.academic_level,
            "status":              p.status.value,
            "progress_percentage": p.progress_percentage,
            "created_at":          p.created_at.isoformat(),
            "completed_at":        p.completed_at.isoformat() if p.completed_at else None,
            "total_marks":         p.result.total_marks  if p.result else None,
            "decision":            p.result.decision     if p.result else None,
        }
        for p in projects
    ]


# ── Detail ─────────────────────────────────────────────────────────────────────

@router.get("/{project_id}", response_model=ProjectDetail)
async def get_project(
    project_id: int,
    user = Depends(get_current_user),
    db   = Depends(get_db_session),
):
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.user_id == user.id,
    ).first()

    if not project:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found")

    from app.core.domain.project import ProjectLifecycle
    lifecycle = ProjectLifecycle(project.status)

    return {
        "id":                  project.id,
        "field_of_study":      project.field_of_study,
        "research_topic":      project.research_topic,
        "academic_level":      project.academic_level,
        "effort_level":        project.effort_level,
        "past_projects_mode":  project.past_projects_mode,
        "status":              project.status.value,
        "progress_percentage": project.progress_percentage,
        "current_phase":       lifecycle.get_phase_description(),
        "created_at":          project.created_at.isoformat(),
        "started_at":          project.started_at.isoformat()   if project.started_at   else None,
        "completed_at":        project.completed_at.isoformat() if project.completed_at else None,
        "has_results":         project.result is not None,
        "total_marks":         project.result.total_marks  if project.result else None,
        "decision":            project.result.decision     if project.result else None,
    }


# ── Delete ─────────────────────────────────────────────────────────────────────

@router.delete("/{project_id}")
async def delete_project(
    project_id: int,
    user = Depends(get_current_user),
    db   = Depends(get_db_session),
):
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.user_id == user.id,
    ).first()

    if not project:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found")

    db.delete(project)
    db.commit()
    return {"success": True, "message": "Project deleted"}


# ── Analytics ──────────────────────────────────────────────────────────────────

@router.get("/{project_id}/analytics")
async def get_project_analytics(
    project_id: int,
    user = Depends(get_current_user),
    db   = Depends(get_db_session),
):
    from app.models.database import ProjectAnalytics

    project = db.query(Project).filter(
        Project.id == project_id,
        Project.user_id == user.id,
    ).first()

    if not project:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found")

    analytics = db.query(ProjectAnalytics).filter(
        ProjectAnalytics.project_id == project_id
    ).first()

    if not analytics:
        return {"message": "Analytics not available yet"}

    return {
        "project_id":              project_id,
        "num_iterations":          analytics.num_iterations,
        "num_web_searches":        analytics.num_web_searches,
        "num_auto_projects_found": analytics.num_auto_projects_found,
        "num_user_projects_analyzed": analytics.num_user_projects_analyzed,
        "final_word_count":        analytics.final_word_count,
        "target_word_count":       analytics.target_word_count,
        "total_generation_time":   analytics.total_generation_time,
        "completeness_score":      analytics.completeness_score,
        "novelty_score":           analytics.novelty_score,
    }


# ── Download ───────────────────────────────────────────────────────────────────
#
# Builds a .docx from the specification_json stored in the database.
# specification_json shape (from research_spec.py / ProjectResult):
# {
#   "project_title": str,
#   "field_of_study": str,
#   "academic_level": str,
#   "total_word_count": int,
#   "abstract":              {"section_name": str, "content": str, "word_count": int},
#   "justification_and_aim": {"section_name": str, "content": str, "word_count": int},
#   "objectives":            {"section_name": str, "content": str, "word_count": int},
#   "literature_review":     {"section_name": str, "content": str, "word_count": int},
#   "methodology":           {"section_name": str, "content": str, "word_count": int},
#   "work_plan":             {"section_name": str, "content": str, "word_count": int},
#   "references":            [str, ...]
# }

@router.get("/{project_id}/download")
async def download_project(
    project_id: int,
    user = Depends(get_current_user),
    db   = Depends(get_db_session),
):
    """
    Download project specification as a formatted .docx file.
    """
    # ── 1. Fetch project ───────────────────────────────────────────────────
    project = db.query(Project).filter(
        Project.id      == project_id,
        Project.user_id == user.id,
    ).first()

    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    if not project.result or not project.result.specification_json:
        raise HTTPException(
            status_code=404,
            detail="Specification not ready yet. Please wait for generation to complete.",
        )

    spec = project.result.specification_json  # already a dict (JSON column)
    review = project.result.final_review_json or {}

    # ── 2. Build .docx ─────────────────────────────────────────────────────
    docx_bytes = _build_docx(spec, review, project)

    # ── 3. Safe filename ───────────────────────────────────────────────────
    title = (spec.get("project_title") or project.research_topic or "specification")
    safe_title = "".join(c if c.isalnum() or c in " -_" else "_" for c in title)[:60].strip()
    filename = f"{safe_title}.docx"

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Content-Length": str(len(docx_bytes)),
        },
    )


# ── DOCX builder ───────────────────────────────────────────────────────────────

def _build_docx(spec: dict, review: dict, project) -> bytes:
    """Build a formatted Word document from specification_json."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Page margins ───────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    PURPLE = RGBColor(0x7C, 0x3A, 0xED)   # Tailwind purple-600
    GRAY   = RGBColor(0x6B, 0x72, 0x80)   # Tailwind gray-500

    def heading1(text: str):
        p = doc.add_heading(text, level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0] if p.runs else p.add_run(text)
        run.font.color.rgb = PURPLE
        run.font.size      = Pt(14)
        run.bold           = True

    def subtext(text: str):
        p   = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size      = Pt(9)
        run.font.color.rgb = GRAY
        run.italic         = True
        p.paragraph_format.space_after = Pt(4)

    def body(text: str):
        if not text or not text.strip():
            doc.add_paragraph("(No content)")
            return
        for para in text.strip().split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    # ── Title page ─────────────────────────────────────────────────────────
    title_para = doc.add_heading(spec.get("project_title", "Research Specification"), level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title_para.runs:
        title_para.runs[0].font.size      = Pt(20)
        title_para.runs[0].font.color.rgb = PURPLE

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"{spec.get('academic_level', '')} · {spec.get('field_of_study', '')}"
        f"   |   {spec.get('total_word_count', 0):,} words"
    ).font.color.rgb = GRAY

    if review:
        marks_para = doc.add_paragraph()
        marks_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        marks_para.add_run(
            f"Score: {review.get('total_marks', '–')}/100  ·  {review.get('decision', '')}"
        ).bold = True

    doc.add_page_break()

    # ── Ordered sections ───────────────────────────────────────────────────
    SECTIONS = [
        ("abstract",             "Abstract"),
        ("justification_and_aim","Justification and Overall Aim"),
        ("objectives",           "Objectives"),
        ("literature_review",    "Review of Literature"),
        ("methodology",          "Methodology"),
        ("work_plan",            "Work Plan"),
    ]

    for key, label in SECTIONS:
        section_data = spec.get(key, {})
        if not section_data:
            continue
        content    = section_data.get("content", "")
        word_count = section_data.get("word_count", len(content.split()))

        heading1(label)
        subtext(f"{word_count:,} words")
        body(content)
        doc.add_paragraph()  # spacing

    # ── References ─────────────────────────────────────────────────────────
    refs = spec.get("references", [])
    if refs:
        heading1("References")
        for i, ref in enumerate(refs, 1):
            p   = doc.add_paragraph(style="List Number")
            p.add_run(ref)

    # ── Review appendix ────────────────────────────────────────────────────
    if review and review.get("total_marks") is not None:
        doc.add_page_break()
        heading1("Professor Review (Appendix)")

        doc.add_paragraph(
            f"Score: {review.get('total_marks')}/100  |  Decision: {review.get('decision', '–')}"
        )

        strengths = review.get("strengths", [])
        if strengths:
            doc.add_heading("Strengths", level=2)
            for s in strengths:
                doc.add_paragraph(s, style="List Bullet")

        issues = review.get("critical_issues", [])
        if issues:
            doc.add_heading("Issues to Address", level=2)
            for issue in issues:
                doc.add_paragraph(issue, style="List Bullet")

        priorities = review.get("improvement_priorities", [])
        if priorities:
            doc.add_heading("Improvement Priorities", level=2)
            for i, p in enumerate(priorities, 1):
                doc.add_paragraph(f"{i}. {p}")

    # ── Serialise to bytes ─────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()