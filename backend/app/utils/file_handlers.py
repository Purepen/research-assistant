"""
File Handling Utilities

Purpose: Safe loading of user uploads (docx, txt, pdf)
"""

from typing import List, Optional
from docx import Document
import os


def load_docx_safely(file_path: str) -> Optional[Document]:
    """
    Safely load a DOCX file
    
    Args:
        file_path: Path to .docx file
        
    Returns:
        Document object or None if failed
    """
    try:
        if not os.path.exists(file_path):
            print(f"❌ File not found: {file_path}")
            return None
        
        doc = Document(file_path)
        return doc
        
    except Exception as e:
        print(f"❌ Error loading {file_path}: {e}")
        return None


def load_user_dumps(file_paths: List[str], max_chars: int = 20000) -> List[tuple[str, str]]:
    """
    Load user-provided project dump files
    
    Args:
        file_paths: List of file paths
        max_chars: Maximum characters to read per file
        
    Returns:
        List of (filename, content) tuples
    """
    loaded = []
    
    for path in file_paths:
        try:
            if not os.path.exists(path):
                print(f"⚠️ Skipping missing file: {path}")
                continue
            
            # Handle different file types
            if path.endswith('.docx'):
                doc = Document(path)
                content = '\n'.join([para.text for para in doc.paragraphs])
            elif path.endswith('.txt'):
                with open(path, 'r', encoding='utf-8') as f:
                    content = f.read()
            else:
                print(f"⚠️ Unsupported file type: {path}")
                continue
            
            # Truncate if too long
            if len(content) > max_chars:
                content = content[:max_chars] + "\n\n[TRUNCATED]"
            
            filename = os.path.basename(path)
            loaded.append((filename, content))
            print(f"   ✓ Loaded: {filename} ({len(content)} chars)")
            
        except Exception as e:
            print(f"❌ Error loading {path}: {e}")
            continue
    
    return loaded
